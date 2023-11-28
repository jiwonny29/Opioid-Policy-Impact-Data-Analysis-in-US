"""
This code make diff in diff to find out if the policies implemented in the states had an effect on the opioid crisis.
"""

import pandas as pd
import numpy as np
from statsmodels.formula.api import ols
from docx import Document
import matplotlib.pyplot as plt
import seaborn as sns
import warnings


time_state = {}
time_state["FL"] = 2010
time_state["WA"] = 2011


data = pd.read_csv("../20_intermediate_files/final_merged_table.csv")
counterfactuals = pd.read_csv(
    "../20_intermediate_files/counterfactuals_by_euclidean_distance.csv"
)


for state in ["FL", "WA"]:
    counterfactuals_by_state = list(
        counterfactuals[counterfactuals["pairing_state"] == state]["State"]
    ) + [state]

    data["morphine_pp"] = data["total_morphine_mg"] / data["total_population"]
    data["time_treatment"] = [1 if x >= time_state[state] else 0 for x in data["year"]]
    data["treated_county"] = [1 if x == state else 0 for x in data["state"]]
    data["treatment"] = data["time_treatment"] * data["treated_county"]

    data_filtered = data[data["state"].isin(counterfactuals_by_state)].copy()

    # linear regression for mprphine bought per person
    ols_morphine = ols(
        "morphine_pp ~ time_treatment + treated_county + treatment", data=data_filtered
    ).fit()

    # linear for mortality rate
    ols_mortality = ols(
        "filled_mortality_rate ~ time_treatment + treated_county + treatment",
        data=data_filtered,
    ).fit()

    # Get  the summary for morphine
    summary_str_morphine = ols_morphine.summary().as_text()

    # Get the summary as a string
    summary_str_mortality = ols_mortality.summary().as_text()

    # Create a new Word document
    doc = Document()

    # Add the summary to the document
    doc.add_paragraph(summary_str_morphine)

    doc.add_paragraph(summary_str_mortality)

    # Save the document
    doc.save("../30_results/" + state + "_regressions.docx")

    ####### PLOTS ##########

    variables = ["morphine_pp", "filled_mortality_rate"]

    for var in variables:
        # plotting morphine per person over time nad its linear regression before treatment
        sns.lmplot(
            x="year",
            y=var,
            hue="treated_county",
            data=data_filtered,
            col="time_treatment",
            scatter=False,
            facet_kws=dict(sharex=False),
        )
        plt.savefig("../30_results/diff_diff_" + state + "_" + var + ".png")

    ######## TEXAS ######### biweekly Analysis

"""
The following code is commented because it 
took the primary Washington post dataset that is
locally saved (100GB).
"""

# from pyspark.sql.functions import date_format, sum, count,to_date, year,dayofmonth, when, month

# # Create a SparkSession
# spark = SparkSession.builder.getOrCreate()


# # As the file was to large I unziped it and saved it in my local machine
# file_path = "../../arcos_all_washpost.tsv"

# # Read the data
# opioid_df = spark.read.csv(file_path, sep="\t", header=True)

# # Convert 'REPORTER_DEA_NO' to date
# opioid_df = opioid_df.withColumn("TRANSACTION_DATE", to_date("TRANSACTION_DATE", "yyyy-MM-dd"))

# # Create 'year' column
# opioid_df = opioid_df.withColumn("year", year("TRANSACTION_DATE"))

# opioid_df = opioid_df.withColumn("month", month("TRANSACTION_DATE"))

# # Create 'day' column based on condition
# opioid_df = opioid_df.withColumn("day", when(dayofmonth("TRANSACTION_DATE") < 15, 1).otherwise(15))

# # Groupby the data

# states_of_interest=['TX','NY','VA','ID']

# #filtering the data by states of interest
# opioid_df_tx=opioid_df.filter(opioid_df.BUYER_STATE.isin(states_of_interest))

# opioid_df_tx = opioid_df_tx.groupBy(
#     "BUYER_STATE", "BUYER_COUNTY", "DRUG_NAME", "year", "month", "day"
# ).agg(
#     sum("MME").alias("total_morphine_mg"),
#     count("REPORTER_DEA_NO").alias("total_transactions"),
# )

# # Converting to pandas
# dataframe_to_print = opioid_df_tx.toPandas()
# dataframe_to_print['date'] = pd.to_datetime(dataframe_to_print[['year', 'month', 'day']])

# # collapsing the data for each biweekly period
# dataframe_to_print=dataframe_to_print['BUYER_STATE BUYER_COUNTY year date total_morphine_mg total_transactions'.split()].groupby('BUYER_STATE BUYER_COUNTY year date'.split(),as_index=False).sum()

# # importing population data
# population = pd.read_csv(
#     "../00_source_data/Additional_datasets/population_by_age_and_sex.csv"
# )
# population=population.rename(columns={'IBRC_Geo_ID': 'fips',
#                            'Year':'year',
#                             'Total Population':'tot_pop',
#                             'Description':'name'
#                            })

# population=population['fips name year tot_pop'.split()]

# fips=pd.read_csv('https://github.com/kjhealy/fips-codes/raw/master/state_and_county_fips_master.csv')

# population=population.merge(fips, on='fips',how='outer',indicator=True)

# population=population[population['_merge']=='both']

# #Preparing counties for merge
# population['COUNTY']=[i[:-7].upper() if 'County' in i else i.upper() for i in population['name_y']]

# population.drop(columns=['_merge','name_y','name_x'],inplace=True)

# #merging population data
# dataframe_to_print=dataframe_to_print.merge(population, left_on=['BUYER_STATE','BUYER_COUNTY','year'],right_on=['state','COUNTY','year'],how='left',indicator=True)

# #creating variables for regression
# dataframe_to_print['morphine_pp']=dataframe_to_print['total_morphine_mg']/dataframe_to_print['tot_pop']
# # I picked 2008 to have 2 years after treatment
# dataframe_to_print['time_treatment']=[1 if x>=2007 else 0 for x in dataframe_to_print['year']]
# dataframe_to_print['treated_county']=[1 if x=='TX' else 0 for x in dataframe_to_print['state']]
# dataframe_to_print['treatment']=dataframe_to_print['time_treatment']*dataframe_to_print['treated_county']

# dataframe_to_print=dataframe_to_print[dataframe_to_print['year']<2009]

# #Saving the data to parquet
# dataframe_to_print[['state', 'COUNTY','year', 'date', 'total_morphine_mg',
#        'total_transactions', 'fips', 'tot_pop',
#        'morphine_pp', 'time_treatment', 'treated_county', 'treatment']].to_parquet('../20_intermediate_files/TX_morphine_biweekly.parquet')

texas_df = pd.read_parquet("../20_intermediate_files/TX_morphine_biweekly.parquet")

ols_morphine = ols(
    "morphine_pp ~ time_treatment + treated_county + treatment", data=texas_df
).fit()

# Get  the summary for morphine
summary_str_morphine = ols_morphine.summary().as_text()

# Create a new Word document
doc = Document()

# Add the summary to the document
doc.add_paragraph(summary_str_morphine)

# Save the document
doc.save("../30_results/TX_regressions.docx")


####### PLOT #######

import matplotlib.dates as mdates

texas_df["date_num"] = mdates.date2num(texas_df["date"])

sns.lmplot(
    x="date_num",
    y="morphine_pp",
    hue="treated_county",
    data=texas_df,
    col="time_treatment",
    scatter=False,
    facet_kws=dict(sharex=False),
)
plt.savefig("../30_results/diff_diff_TX_morphine_pp.png")
